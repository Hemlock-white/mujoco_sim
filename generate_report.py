#!/usr/bin/env python3
"""
SDK2 Architecture Analysis Report Generator
Generates a professional DOCX report for the mujoco_sim SDK2 communication system.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ─── Page setup ────────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.0)

# ─── Helper styles ─────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
    return p

def add_para(doc, text, bold=False, italic=False, size=10, indent=0):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p

def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x00, 0x60, 0x00)
    return p

def add_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run("⚠  " + text)
    run.font.size  = Pt(9)
    run.font.color.rgb = RGBColor(0xC0, 0x50, 0x00)
    run.bold = True
    return p

def table_header_row(table, headers, bg="1F3864"):
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        cell.text = h
        set_cell_bg(cell, bg)
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.bold  = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size  = Pt(9)

def add_table_row(table, values, bg=None):
    row = table.add_row()
    for i, v in enumerate(values):
        cell = row.cells[i]
        cell.text = str(v)
        if bg:
            set_cell_bg(cell, bg)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(9)
    return row

# ══════════════════════════════════════════════════════════════════════════════
#  TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run("MuJoCo-Sim SDK2 通訊架構\n技術分析報告書")
title_run.font.size  = Pt(22)
title_run.font.bold  = True
title_run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

doc.add_paragraph()
sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = sub_p.add_run("Unitree Go2 四足機器人 · Convex-MPC · DDS 通訊")
sub_run.font.size  = Pt(13)
sub_run.font.color.rgb = RGBColor(0x40, 0x40, 0x80)

doc.add_paragraph()
doc.add_paragraph()

meta_p = doc.add_paragraph()
meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta_run = meta_p.add_run(
    "Branch: pd_sit2stand   |   Repo: mujoco_sim\n"
    "分析日期：2026-05-10   |   作者：系統自動生成"
)
meta_run.font.size = Pt(10)
meta_run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 1 – OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "1. 系統概述", 1, (0x1F, 0x38, 0x64))

add_para(doc,
    "本報告針對 mujoco_sim 專案中所有 SDK2 相關程式進行深度技術分析，"
    "涵蓋 Teleop 遙控端、MPC 控制器（mpc_locomotion_sdk2.py）、"
    "MuJoCo 仿真橋接器（mujoco_sim_sdk2.py + UnitreeSdk2Bridge）"
    "三大核心模組，並詳細說明其間透過 Unitree DDS（CycloneDDS）"
    "進行的低延遲通訊機制、執行緒配置、訊號頻率及潛在問題。",
    size=10)

doc.add_paragraph()

# ─── 1.1 Scope table ───────────────────────────────────────────────────────
add_heading(doc, "1.1 SDK2 相關程式清單", 2)

scope_table = doc.add_table(rows=1, cols=4)
scope_table.style = 'Table Grid'
scope_table.alignment = WD_TABLE_ALIGNMENT.CENTER
table_header_row(scope_table, ["程式 / 模組", "路徑", "角色", "執行方式"], "1F3864")

scope_rows = [
    ("teleop_client.py",            "專案根目錄",                   "遙控指令產生器",              "手動執行 python3 teleop_client.py"),
    ("mujoco_sim/udp_reader.py",    "mujoco_sim/",                  "UDP 封包接收（模擬 Gamepad）", "由 mpc_locomotion_sdk2.py 匯入"),
    ("mujoco_sim/config_sdk2.py",   "mujoco_sim/",                  "全域參數設定",                 "由各模組匯入"),
    ("mpc_locomotion_sdk2.py",      "專案根目錄",                   "MPC 控制主程式",               "手動執行 python3 mpc_locomotion_sdk2.py"),
    ("mujoco_sim/mujoco_sim_utils.py","mujoco_sim/",                "感測器資料解析 / PD 控制函式", "由 mpc_locomotion_sdk2.py 匯入"),
    ("mujoco_sim/unitree_sdk2py_bridge.py","mujoco_sim/",           "DDS ↔ MuJoCo 橋接器",          "由 mujoco_sim_sdk2.py 匯入"),
    ("mujoco_sim_sdk2.py",          "專案根目錄",                   "MuJoCo 仿真 + Viewer 主程式",  "手動執行 python3 mujoco_sim_sdk2.py"),
    ("assets/go2/scene.xml",        "assets/go2/",                  "GO2 MuJoCo 場景定義",          "由 mujoco_sim_sdk2.py 載入"),
    ("MPC_Controller/",             "MPC_Controller/",              "完整 MPC 控制器函式庫",         "由 mpc_locomotion_sdk2.py 匯入"),
]

for i, row_data in enumerate(scope_rows):
    bg = "EBF3FF" if i % 2 == 0 else "FFFFFF"
    add_table_row(scope_table, row_data, bg)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 2 – ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "2. 整體架構說明", 1, (0x1F, 0x38, 0x64))

add_para(doc,
    "SDK2 系統由三個獨立 Process（進程）組成，透過 UDP（本機）與 DDS 兩種通訊協定互連，"
    "形成一個閉迴路控制系統（Closed-loop control pipeline）。",
    size=10)

add_para(doc,
    "Process A：teleop_client.py  →（UDP 127.0.0.1:9876）→  "
    "Process B：mpc_locomotion_sdk2.py  →（DDS rt/lowcmd）→  "
    "Process C：mujoco_sim_sdk2.py  →（DDS rt/lowstate）→  回到 Process B",
    bold=True, size=10, indent=0.5)

doc.add_paragraph()

# ─── 2.1 Process descriptions ─────────────────────────────────────────────
add_heading(doc, "2.1 Process A：Teleop 遙控端（teleop_client.py）", 2)

add_para(doc,
    "負責產生人機互動介面，讓使用者透過鍵盤輸入控制指令並以 UDP 封包傳送至 MPC 控制器。"
    "程式以 JSON 格式封裝速度指令（vx, vy, wz）、模式切換（mode）及緊急停止（estop）"
    "欄位，經 socket 發至 127.0.0.1:9876。",
    size=10)

add_para(doc, "主要流程：", bold=True, size=10)
steps_A = [
    "1. main() 啟動，顯示操作說明選單。",
    "2. 使用者輸入 'move' → 進入 wsad() 副迴圈，接受 WASD/QE/Z/C 鍵控制。",
    "   每次按鍵後立即呼叫 send_state()，透過 UDP 發送目前 state 字典。",
    "3. 使用者輸入 'stand' → state['mode']='stand'，立即發送。",
    "4. 使用者輸入 'estop' → state['estop']=True，立即發送。",
    "5. 使用者輸入 'clear' → 解除 estop，發送後重置 clear_estop flag。",
    "6. 使用者輸入 'exit' → 離開主迴圈，程式結束。",
]
for s in steps_A:
    add_para(doc, s, size=9.5, indent=0.5)

add_note(doc,
    "【潛在問題】心跳包（heartbeat）執行緒在程式碼中已被注解掉（#heart_thread.start()），"
    "表示若使用者無操作則不會定期傳送封包。"
    "udp_reader 端有 50ms socket timeout，長時間無封包不影響穩定，但無自動重連機制。")

doc.add_paragraph()

add_heading(doc, "2.2 Process B：MPC 控制器（mpc_locomotion_sdk2.py）", 2)

add_para(doc,
    "核心控制程式，負責：(1) 接收 Teleop UDP 指令；(2) 訂閱 DDS rt/lowstate（仿真狀態）；"
    "(3) 計算 Convex-MPC 力矩或 PD 站立力矩；(4) 發佈 DDS rt/lowcmd（馬達控制指令）。",
    size=10)

add_para(doc, "類別 MPCLocomotionSDK2 結構：", bold=True, size=10)

class_table = doc.add_table(rows=1, cols=3)
class_table.style = 'Table Grid'
class_table.alignment = WD_TABLE_ALIGNMENT.CENTER
table_header_row(class_table, ["方法 / 屬性", "型別", "功能說明"], "2E5090")

class_rows = [
    ("__init__()",          "Constructor",       "初始化 low_cmd、low_state=None、CRC 計算器"),
    ("Init()",              "方法",              "建立 DDS Publisher(rt/lowcmd) + Subscriber(rt/lowstate)；呼叫 InitLowCmd()"),
    ("InitLowCmd()",        "私有方法",          "設定 low_cmd 標頭 (0xFE,0xEF)、level_flag=0xFF；20 顆馬達初始化為阻尼模式"),
    ("LowStateMessageHandler()", "DDS Callback", "接收 rt/lowstate 後存入 self.low_state"),
    ("Start()",             "方法",              "啟動 RecurrentThread(interval=0.002s)，呼叫 LowCmdWrite"),
    ("LowCmdWrite()",       "執行緒目標函式",    "計算 CRC 後發佈 rt/lowcmd，頻率 500 Hz"),
    ("MPC_RUN()",           "主控制迴圈",        "200 Hz 迴圈；依 gamepad 狀態切換 pd_stand_sdk2 / MPC 計算；寫入 motor_cmd.tau"),
    ("low_cmd",             "LowCmd_",           "DDS 馬達控制指令訊息"),
    ("low_state",           "LowState_",         "最新 DDS 仿真狀態（None 表示尚未收到）"),
    ("lowCmdWriteThreadPtr","RecurrentThread",   "500 Hz 發佈執行緒"),
]

for i, row_data in enumerate(class_rows):
    bg = "EBF3FF" if i % 2 == 0 else "FFFFFF"
    add_table_row(class_table, row_data, bg)

doc.add_paragraph()

add_para(doc, "MPC_RUN() 主迴圈邏輯（200 Hz，sleep=0.005s）：", bold=True, size=10)
add_para(doc, "① 若 low_state 為 None → 等待 rt/lowstate 首包（列印提示，不阻塞執行緒）", size=9.5, indent=0.5)
add_para(doc, "② gamepad.is_standing == True → 呼叫 pd_stand_sdk2() 計算 12 顆馬達力矩", size=9.5, indent=0.5)
add_para(doc, "③ gamepad.is_moving == True：", size=9.5, indent=0.5)
add_para(doc, "   a. 取 UDP 速度指令 (vx, vy, wz) 及 gait/mode", size=9.5, indent=0.8)
add_para(doc, "   b. get_dof_state_sdk2() 解析關節角度/速度（R↔L 軸序交換）", size=9.5, indent=0.8)
add_para(doc, "   c. get_body_state_sdk2() 解析 IMU 四元數/角速度/加速度", size=9.5, indent=0.8)
add_para(doc, "   d. robotRunner.run() → 返回 12D float32 力矩向量", size=9.5, indent=0.8)
add_para(doc, "④ 將力矩寫入 low_cmd.motor_cmd[i].tau（注意 R↔L 順序再次交換）", size=9.5, indent=0.5)
add_para(doc, "⑤ 若 Parameters.locomotionUnsafe → 觸發 fake_event 清除 estop flag", size=9.5, indent=0.5)

add_note(doc,
    "【潛在問題】MPC_RUN() 在 use_gamepad==False 時立即 break，等同程式無任何控制輸出。"
    "此設計為強制要求 UDP Gamepad，若 --disable-gamepad 被傳入則整個控制迴圈空轉。")

add_note(doc,
    "【潛在問題】MPC_RUN() 200Hz 與 LowCmdWrite() 500Hz 共用 self.low_cmd，"
    "兩個執行緒間並無 Lock 保護，存在 data race 風險，"
    "可能在 CRC 計算的同時 MPC_RUN 正在寫入 tau 值。")

doc.add_paragraph()

add_heading(doc, "2.3 Process C：MuJoCo 仿真端（mujoco_sim_sdk2.py）", 2)

add_para(doc,
    "負責載入 GO2 MJCF 模型、執行物理仿真步進、管理 Viewer 渲染，"
    "並透過 UnitreeSdk2Bridge 與 DDS 通訊。",
    size=10)

add_para(doc, "初始化序列（在 __main__ 之前的模組頂層執行）：", bold=True, size=10)
init_steps = [
    "1. 載入 assets/go2/scene.xml → mj_model, mj_data",
    "2. 以 mujoco.viewer.launch_passive() 建立被動 Viewer（不佔用主執行緒）",
    "3. 設定攝影機（lookat=[0,0,0.25]、distance=2.0、azimuth=90、elevation=-20）",
    "4. 設定 mj_model.opt.timestep = config.SIMULATE_DT（預設 0.005s = 200Hz）",
    "5. time.sleep(0.2) 等待 Viewer 初始化",
]
for s in init_steps:
    add_para(doc, s, size=9.5, indent=0.5)

add_para(doc, "執行緒結構：", bold=True, size=10)

threads_table = doc.add_table(rows=1, cols=5)
threads_table.style = 'Table Grid'
threads_table.alignment = WD_TABLE_ALIGNMENT.CENTER
table_header_row(threads_table, ["執行緒名稱", "目標函式", "週期", "頻率", "功能"], "2E5090")

threads_data = [
    ("sim_thread",          "SimulationThread()",      "0.005s",  "200 Hz", "DDS 初始化、mj_step() 物理步進"),
    ("viewer_thread",       "PhysicsViewerThread()",   "0.02s",   "50 Hz",  "viewer.sync() 畫面更新"),
    ("sim_lowstate",        "PublishLowState()",        "0.005s",  "200 Hz", "DDS 發佈 rt/lowstate"),
    ("sim_highstate",       "PublishHighState()",       "0.005s",  "200 Hz", "DDS 發佈 rt/sportmodestate"),
    ("writebasiccmd",       "LowCmdWrite()",            "0.002s",  "500 Hz", "DDS 發佈 rt/lowcmd（於 Process B）"),
]
for i, row_data in enumerate(threads_data):
    bg = "EBF3FF" if i % 2 == 0 else "FFFFFF"
    add_table_row(threads_table, row_data, bg)

doc.add_paragraph()

add_note(doc,
    "【潛在問題】sim_thread 與 viewer_thread 共用 mj_data，"
    "透過 threading.Lock()（locker）保護，但 mj_step 與 viewer.sync 皆需 lock，"
    "當 viewer.sync() 執行時間超過 0.005s 時 sim_thread 會被迫等待，"
    "導致仿真速度低於理論 200Hz。")

add_note(doc,
    "【潛在問題】UnitreeSdk2Bridge 於 SimulationThread 內部建立，"
    "其 PublishLowState / PublishHighState 執行緒以 self.dt（= SIMULATE_DT=0.005）為間隔，"
    "但仿真步進也在 SimulationThread 的迴圈中。理論上 PublishLowState 與 mj_step 的資料"
    "可能有 race condition（mj_data.sensordata 在 mj_step 執行中），"
    "目前只有 locker 保護 mj_step 與 viewer.sync，橋接器讀取 sensordata 並未在 lock 內。")

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 3 – DDS COMMUNICATION
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "3. DDS 通訊詳解", 1, (0x1F, 0x38, 0x64))

add_para(doc,
    "系統使用 Unitree SDK2（基於 CycloneDDS）進行進程間通訊。"
    "所有 DDS 節點初始化使用 DOMAIN_ID=1、INTERFACE=eth0。",
    size=10)

add_note(doc,
    "【重要】config_sdk2.py 設定 INTERFACE='eth0'，"
    "代表 DDS 通訊綁定至 eth0 網路介面。"
    "原始碼註解說明此機器無 multicast 支援，因此改用 eth0 而非預設介面。"
    "在 WSL2 環境下，eth0 為 WSL 內部介面，兩個 Process 皆在同機運行時可正常通訊；"
    "若要跨機通訊需確認 eth0 路由設定。")

doc.add_paragraph()

add_heading(doc, "3.1 DDS Topics 完整列表", 2)

dds_table = doc.add_table(rows=1, cols=6)
dds_table.style = 'Table Grid'
dds_table.alignment = WD_TABLE_ALIGNMENT.CENTER
table_header_row(dds_table,
    ["Topic 名稱", "IDL 型別", "Publisher", "Subscriber", "頻率", "資料內容"],
    "1F3864")

dds_rows = [
    ("rt/lowcmd",
     "LowCmd_\n(unitree_go)",
     "mpc_locomotion_sdk2.py\n(MPCLocomotionSDK2.LowCmdWrite)\n500 Hz RecurrentThread",
     "mujoco_sim_sdk2.py\n(UnitreeSdk2Bridge.LowCmdHandler)\n事件觸發",
     "500 Hz",
     "head[2]、level_flag、gpio\n20×MotorCmd：mode,q,kp,dq,kd,tau\ncrc（CRC32）"),
    ("rt/lowstate",
     "LowState_\n(unitree_go)",
     "mujoco_sim_sdk2.py\n(UnitreeSdk2Bridge.PublishLowState)\n200 Hz RecurrentThread",
     "mpc_locomotion_sdk2.py\n(MPCLocomotionSDK2.LowStateMessageHandler)\n事件觸發",
     "200 Hz",
     "12×MotorState：q,dq,tau_est\nIMU：quaternion[4],gyroscope[3]\n       accelerometer[3]\npower_v, power_a"),
    ("rt/sportmodestate",
     "SportModeState_\n(unitree_go)",
     "mujoco_sim_sdk2.py\n(UnitreeSdk2Bridge.PublishHighState)\n200 Hz RecurrentThread",
     "目前無訂閱者\n（保留供擴充）",
     "200 Hz",
     "position[3]（感測器 idx 46-48）\nvelocity[3]（感測器 idx 49-51）"),
]

for i, row_data in enumerate(dds_rows):
    bg = "EBF3FF" if i % 2 == 0 else "FFFFFF"
    add_table_row(dds_table, row_data, bg)  # placeholder, rewritten below

# rebuild properly
dds_table2 = doc.add_table(rows=1, cols=6)
dds_table2.style = 'Table Grid'
dds_table2.alignment = WD_TABLE_ALIGNMENT.CENTER
# remove placeholder
dds_table._element.getparent().remove(dds_table._element)

table_header_row(dds_table2,
    ["Topic 名稱", "IDL 型別", "Publisher（程式/執行緒）", "Subscriber（程式）", "頻率", "主要資料欄位"],
    "1F3864")

real_dds_rows = [
    ("rt/lowcmd",
     "LowCmd_",
     "mpc_locomotion_sdk2\nwritebasiccmd thread",
     "mujoco_sim_sdk2\nLowCmdHandler (callback)",
     "500 Hz",
     "motor_cmd[0-19]\n.tau .q .dq .kp .kd .mode\nhead crc"),
    ("rt/lowstate",
     "LowState_",
     "mujoco_sim_sdk2\nsim_lowstate thread",
     "mpc_locomotion_sdk2\nLowStateMessageHandler (callback)",
     "200 Hz",
     "motor_state[0-11]\n.q .dq .tau_est\nimu_state.quaternion[4]\n.gyroscope[3]\n.accelerometer[3]"),
    ("rt/sportmodestate",
     "SportModeState_",
     "mujoco_sim_sdk2\nsim_highstate thread",
     "（目前無）",
     "200 Hz",
     "position[3]\nvelocity[3]"),
]

for i, row_data in enumerate(real_dds_rows):
    bg = "EBF3FF" if i % 2 == 0 else "FFFFFF"
    add_table_row(dds_table2, row_data, bg)

doc.add_paragraph()

# ─── 3.2 Motor index mapping ─────────────────────────────────────────────
add_heading(doc, "3.2 馬達索引對應關係", 2)

add_para(doc,
    "SDK2 與 MPC Controller 之間存在左右腿順序差異（R↔L swap）。"
    "下表說明 IDL motor_cmd/motor_state 索引與 MPC 腿部順序的對應關係。",
    size=10)

motor_table = doc.add_table(rows=1, cols=5)
motor_table.style = 'Table Grid'
motor_table.alignment = WD_TABLE_ALIGNMENT.CENTER
table_header_row(motor_table,
    ["SDK2 motor_cmd 索引", "物理關節名稱", "MPC 腿索引", "MPC 陣列索引", "交換規則"],
    "2E5090")

motor_rows = [
    ("0", "FR_0 (Right Front Hip)",  "FL leg 0", "3",  "i%6<=2: SDK[i]→MPC[i+3]"),
    ("1", "FR_1 (Right Front Thigh)","FL leg 0", "4",  ""),
    ("2", "FR_2 (Right Front Knee)", "FL leg 0", "5",  ""),
    ("3", "FL_0 (Left Front Hip)",   "FR leg 0", "0",  "i%6>2: SDK[i]→MPC[i-3]"),
    ("4", "FL_1 (Left Front Thigh)", "FR leg 0", "1",  ""),
    ("5", "FL_2 (Left Front Knee)",  "FR leg 0", "2",  ""),
    ("6", "RR_0 (Right Rear Hip)",   "RL leg 2", "9",  "i%6<=2: SDK[i]→MPC[i+3]"),
    ("7", "RR_1 (Right Rear Thigh)", "RL leg 2", "10", ""),
    ("8", "RR_2 (Right Rear Knee)",  "RL leg 2", "11", ""),
    ("9", "RL_0 (Left Rear Hip)",    "RR leg 2", "6",  "i%6>2: SDK[i]→MPC[i-3]"),
    ("10","RL_1 (Left Rear Thigh)",  "RR leg 2", "7",  ""),
    ("11","RL_2 (Left Rear Knee)",   "RR leg 2", "8",  ""),
]

for i, row_data in enumerate(motor_rows):
    bg = "EBF3FF" if i % 2 == 0 else "FFFFFF"
    add_table_row(motor_table, row_data, bg)

add_note(doc,
    "【潛在問題】此 R↔L 軸序交換在 get_dof_state_sdk2()（讀取）與 MPC_RUN()（寫入）"
    "兩處各自實作一次相同邏輯，若任一處更新不一致，"
    "將導致力矩施加到錯誤的馬達，造成機器人動作異常。"
    "建議將交換邏輯抽取為單一工具函式（如 sdk2_to_mpc_index()）集中管理。")

doc.add_paragraph()

# ─── 3.3 Sensor index mapping ─────────────────────────────────────────────
add_heading(doc, "3.3 MuJoCo 感測器索引對應（UnitreeSdk2Bridge）", 2)

add_para(doc,
    "UnitreeSdk2Bridge.PublishLowState() 依 sensordata 陣列偏移量讀取仿真狀態，"
    "對應關係如下（num_motor = 12）：",
    size=10)

sensor_table = doc.add_table(rows=1, cols=4)
sensor_table.style = 'Table Grid'
sensor_table.alignment = WD_TABLE_ALIGNMENT.CENTER
table_header_row(sensor_table,
    ["sensordata 索引", "資料意義", "對應 LowState 欄位", "備註"],
    "2E5090")

sensor_rows = [
    ("0 – 11",     "12 顆馬達關節角度 q",      "motor_state[0-11].q",          "MOTOR_SENSOR_NUM=3"),
    ("12 – 23",    "12 顆馬達關節速度 dq",      "motor_state[0-11].dq",         ""),
    ("24 – 35",    "12 顆馬達估計扭矩 tau_est", "motor_state[0-11].tau_est",    ""),
    ("36 – 39",    "IMU 四元數 (w,x,y,z)",       "imu_state.quaternion[0-3]",    "have_frame_sensor_ 旗標需為 True"),
    ("40 – 42",    "IMU 角速度 (x,y,z)",          "imu_state.gyroscope[0-2]",     ""),
    ("43 – 45",    "IMU 線加速度 (x,y,z)",        "imu_state.accelerometer[0-2]", ""),
    ("46 – 48",    "機體位置 (x,y,z)",            "high_state.position[0-2]",     "PublishHighState"),
    ("49 – 51",    "機體速度 (x,y,z)",            "high_state.velocity[0-2]",     "PublishHighState"),
]

for i, row_data in enumerate(sensor_rows):
    bg = "EBF3FF" if i % 2 == 0 else "FFFFFF"
    add_table_row(sensor_table, row_data, bg)

add_note(doc,
    "【潛在問題】PublishLowState() 以 have_frame_sensor_ 而非 have_frame_sensor"
    "（缺少底線）作為 IMU 資料讀取的判斷條件，"
    "但在 __init__ 中初始化的是 self.have_frame_sensor（無底線）與 self.have_imu（無底線）。"
    "屬性 have_frame_sensor_ 及 have_imu_ 只在迴圈內條件成立時才被 set，"
    "若感測器名稱不符合條件，這兩個屬性永遠不會被建立，"
    "呼叫時將拋出 AttributeError。"
    "且讀取 IMU 的條件錯誤地檢查 have_frame_sensor_ 而非 have_imu_。")

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 4 – COMMUNICATION FLOWCHART (TEXT-BASED)
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "4. 通訊 Flowchart（文字圖）", 1, (0x1F, 0x38, 0x64))

add_para(doc,
    "以下流程圖說明從使用者輸入到 MuJoCo 物理仿真更新的完整資料流。",
    size=10)

flowchart_text = """
┌─────────────────────────────────────────────────────────────────────────────────┐
│                      整體閉迴路控制流程（Closed-loop）                          │
└─────────────────────────────────────────────────────────────────────────────────┘

  ┌──────────────────────────┐
  │   使用者鍵盤輸入          │
  │  (teleop_client.py)      │
  │  WSAD/QE/Z/模式切換/estop │
  └────────────┬─────────────┘
               │ UDP 127.0.0.1:9876
               │ JSON封包 {vx,vy,wz,mode,estop}
               │ 非定期（按鍵觸發）
               ▼
  ┌──────────────────────────────────────────┐
  │        Process B: mpc_locomotion_sdk2.py  │
  │  ┌──────────────────────────────────────┐ │
  │  │ UDPGamepad._listen() [daemon thread] │ │
  │  │  ← socket.recvfrom(1024)            │ │
  │  │  更新 vx,vy,wz,_mode,_estop         │ │
  │  └──────────────────────────────────────┘ │
  │                                            │
  │  ┌──────────────────────────────────────┐ │
  │  │ MPC_RUN() [主執行緒, 200Hz]          │ │
  │  │  ← low_state (DDS rt/lowstate)       │ │
  │  │                                      │ │
  │  │  if is_standing:                     │ │
  │  │    pd_stand_sdk2() → τ(12D)          │ │
  │  │  if is_moving:                       │ │
  │  │    get_dof_state_sdk2()              │ │
  │  │    get_body_state_sdk2()             │ │
  │  │    RobotRunnerFSM.run() → τ(12D)     │ │
  │  │  write τ → low_cmd.motor_cmd[i].tau │ │
  │  └──────────────────────────────────────┘ │
  │                                            │
  │  ┌──────────────────────────────────────┐ │
  │  │ LowCmdWrite() [RecurrentThread,500Hz]│ │
  │  │  CRC() → publish rt/lowcmd           │ │
  │  └──────────────────────────────────────┘ │
  └────────────┬─────────────────────────────┘
               │ DDS rt/lowcmd [LowCmd_]
               │ 500 Hz
               │ motor_cmd[0-19].tau/kp/kd/q/dq
               ▼
  ┌──────────────────────────────────────────────┐
  │   Process C: mujoco_sim_sdk2.py               │
  │   UnitreeSdk2Bridge                           │
  │  ┌────────────────────────────────────────┐  │
  │  │ LowCmdHandler() [DDS callback]         │  │
  │  │  mj_data.ctrl[i] =                     │  │
  │  │   tau + kp*(q_des - q) + kd*(dq_des-dq)│  │
  │  └────────────────────────────────────────┘  │
  │                                               │
  │  ┌────────────────────────────────────────┐  │
  │  │ SimulationThread() [sim_thread, 200Hz] │  │
  │  │  locker.acquire()                      │  │
  │  │  mujoco.mj_step(mj_model, mj_data)     │  │
  │  │  locker.release()                      │  │
  │  └────────────────────────────────────────┘  │
  │                                               │
  │  ┌────────────────────────────────────────┐  │
  │  │ PhysicsViewerThread()[viewer_thread,50Hz│  │
  │  │  locker.acquire()                      │  │
  │  │  viewer.sync()  ← 渲染到螢幕           │  │
  │  │  locker.release()                      │  │
  │  └────────────────────────────────────────┘  │
  │                                               │
  │  ┌────────────────────────────────────────┐  │
  │  │ PublishLowState()[sim_lowstate,200Hz]  │  │
  │  │  → DDS rt/lowstate                    │  │
  │  └────────────────────────────────────────┘  │
  └────────────┬──────────────────────────────────┘
               │ DDS rt/lowstate [LowState_]
               │ 200 Hz
               │ motor_state[0-11].q/dq/tau_est
               │ imu_state.quaternion/gyro/accel
               └──────────────────────────────────→ 回到 Process B MPC_RUN()
"""

p = doc.add_paragraph()
run = p.add_run(flowchart_text)
run.font.name = 'Courier New'
run.font.size = Pt(7.5)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 5 – THREAD & TOPIC DIAGRAM
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "5. 執行緒與 Topic 通訊圖", 1, (0x1F, 0x38, 0x64))

add_para(doc,
    "下圖以文字形式呈現所有執行緒節點（方框）、通訊 Topic（箭頭）及訊號頻率（標註）。"
    "實線箭頭（──→）代表 DDS 訊息；虛線箭頭（--→）代表 UDP 封包；"
    "點線（···→）代表直接函式呼叫或記憶體存取。",
    size=10)

thread_diagram = """
════════════════════════════════════════════════════════════════════════════════════
  PROCESS A: teleop_client.py
  ─────────────────────────────────────────────────────────────────────────────────
  [ main() / wsad() ]  ──UDP 127.0.0.1:9876 JSON~~→  [ UDPGamepad._listen() ]
  非定期 / 按鍵觸發                                   daemon thread (blocking recv)
════════════════════════════════════════════════════════════════════════════════════
  PROCESS B: mpc_locomotion_sdk2.py
  ─────────────────────────────────────────────────────────────────────────────────
                                    ┌─────────────────────────────────────────┐
  [ UDPGamepad._listen() ] ···→     │  MPC_RUN()  [主執行緒  200 Hz]          │
  (daemon, UDP recv loop)           │  讀 low_state / 寫 low_cmd.motor_cmd.tau│
                                    └──────────────────┬──────────────────────┘
                                                       │ 共用 self.low_cmd（無鎖）
                                    ┌──────────────────▼──────────────────────┐
                                    │  LowCmdWrite() [writebasiccmd  500 Hz]  │
                                    │  CRC → ChannelPublisher.Write(low_cmd)  │
                                    └──────────────────┬──────────────────────┘
                                                       │
                                        DDS rt/lowcmd  │  500 Hz
                                         ──────────────▼──────────────
════════════════════════════════════════════════════════════════════════════════════
  PROCESS C: mujoco_sim_sdk2.py
  ─────────────────────────────────────────────────────────────────────────────────
                        ┌────────────────────────────────────────────────────────┐
  DDS rt/lowcmd ──────→ │ LowCmdHandler() [DDS callback, 事件觸發]              │
  500 Hz                │ mj_data.ctrl[i] = τ + kp*(q_des-q) + kd*(dq_des-dq)  │
                        └───────────────────────────────┬────────────────────────┘
                                                        │ 寫入 mj_data.ctrl
                        ┌───────────────────────────────▼────────────────────────┐
                        │ SimulationThread() [sim_thread  200 Hz]                │
                        │ locker.acquire()                                       │
                        │ mj_step(mj_model, mj_data)   ← 物理積分               │
                        │ locker.release()                                       │
                        └─────────┬──────────────────────────────────────────────┘
                                  │ mj_data.sensordata 更新
                    ┌─────────────▼──────────────────┐   ┌─────────────────────────┐
                    │ sim_lowstate [200 Hz]           │   │ sim_highstate [200 Hz]  │
                    │ PublishLowState()               │   │ PublishHighState()      │
                    │  → rt/lowstate                  │   │  → rt/sportmodestate   │
                    └────────────────────────────────┘   └─────────────────────────┘
                                  │
                    ┌─────────────▼──────────────────┐
                    │ PhysicsViewerThread() [50 Hz]  │
                    │ locker.acquire()               │
                    │ viewer.sync() ← 渲染畫面       │
                    │ locker.release()               │
                    └────────────────────────────────┘

     DDS rt/lowstate (200 Hz) ──────────────────────────────────────────────────→
                                                         Process B MPC_RUN() 更新
════════════════════════════════════════════════════════════════════════════════════
"""

p2 = doc.add_paragraph()
run2 = p2.add_run(thread_diagram)
run2.font.name = 'Courier New'
run2.font.size = Pt(7.5)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 6 – MPC CONTROLLER INTERNALS
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "6. MPC 控制器內部架構", 1, (0x1F, 0x38, 0x64))

add_para(doc,
    "RobotRunnerFSM.run() 為 MPC_RUN() 在 is_moving 模式下的核心呼叫，"
    "每次呼叫在主執行緒執行，並隱性地以 200Hz 頻率被 MPC_RUN() 觸發。",
    size=10)

add_heading(doc, "6.1 RobotRunnerFSM.run() 呼叫鏈", 2)

callchain = """
  RobotRunnerFSM.run(dof_states, body_states, commands)
   ├─ DesiredStateCommand.updateCommand(commands)
   │    └─ 更新 x_vel_cmd, y_vel_cmd, yaw_turn_rate
   │
   ├─ LegController.updateData(dof_states)
   │    └─ 12D 關節 q, qd → 4×LegControllerData
   │         └─ computeLegJacobianAndPosition(leg)  ← 解析運動學 J, p, v
   │
   ├─ LegController.zeroCommand()
   │    └─ 清除上一步力矩指令，防止殘留
   │
   ├─ StateEstimator.update(body_states)
   │    └─ 解析 IMU 四元數 → rBody, rpy, vBody, omegaBody
   │         └─ 計算 ground_R_body_frame
   │
   ├─ ControlFSM.runFSM()
   │    ├─ NORMAL 模式：
   │    │    checkTransition() → 無切換 → currentState.run()
   │    │
   │    ├─ TRANSITIONING 模式：
   │    │    currentState.transition() → transitionDone=True
   │    │    → currentState.onExit() → nextState.onEnter()
   │    │    → 回到 NORMAL
   │    │
   │    └─ 狀態清單：
   │         PASSIVE(0), LOCOMOTION(4), RECOVERY_STAND(6)
   │
   │    若 LOCOMOTION：
   │    └─ FSM_State_Locomotion.run()
   │         └─ ConvexMPCLocomotion.run(_data)
   │              └─ 求解 QP 問題（每 iterationsBetweenMPC 步執行一次）
   │                   iterationsBetweenMPC = 27/(1000*0.01) = 2.7 ≈ 2
   │                   dtMPC = 0.01 * 2.7 = 0.027s ≈ 37 Hz
   │
   │    若 RECOVERY_STAND：
   │    └─ FSM_State_RecoveryStand.run()
   │         └─ _FoldLegs() → _StandUp() → _RollOver()
   │              └─ 線性插值關節位置 PD 控制
   │
   └─ LegController.updateCommand()
        └─ 計算最終力矩：
             τ = τ_ff + J^T(f_ff + Kp*(p_des-p) + Kd*(v_des-v))
                       + Kp_joint*(q_des-q) + Kd_joint*(qd_des-qd)
        └─ 返回 legTorques (12,) float32
"""

p3 = doc.add_paragraph()
run3 = p3.add_run(callchain)
run3.font.name = 'Courier New'
run3.font.size = Pt(8)

doc.add_paragraph()

add_heading(doc, "6.2 Convex MPC 求解頻率", 2)
add_para(doc,
    "ConvexMPCLocomotion 於初始化時設定：",
    size=10)
add_code_block(doc,
    "self.cMPC = ConvexMPCLocomotion(\n"
    "    Parameters.controller_dt,          # dt = 0.01s\n"
    "    27/(1000.0*Parameters.controller_dt)  # iterations = 27/(1000*0.01) = 2.7\n"
    ")")
add_para(doc,
    "MPC 每 iterationsBetweenMPC（=2 次 FSM 迭代）求解一次 QP，"
    "對應實際頻率 ≈ 1/(0.01×2) = 50 Hz。"
    "MPC horizon 為 10 步，每步 dtMPC=0.027s，預測範圍約 0.27s。",
    size=10)

add_note(doc,
    "【潛在問題】Parameters.controller_dt = 0.01s，但 MPC_RUN() 的 sleep(0.005) "
    "使主迴圈實際以 200Hz 執行。"
    "RobotRunnerFSM.run() 每次呼叫後都觸發 ControlFSM，"
    "意味著 FSM 以 200Hz 而非 controller_dt 所預期的 100Hz 執行。"
    "ConvexMPCLocomotion 的 iterationsBetweenMPC 計算基於 controller_dt=0.01，"
    "若實際迴圈為 0.005s，MPC 有效頻率將是預設值的兩倍（~100Hz），可能影響控制穩定性。")

doc.add_paragraph()

# ─── 6.3 FSM States table ──────────────────────────────────────────────────
add_heading(doc, "6.3 FSM 狀態機說明", 2)

fsm_table = doc.add_table(rows=1, cols=5)
fsm_table.style = 'Table Grid'
fsm_table.alignment = WD_TABLE_ALIGNMENT.CENTER
table_header_row(fsm_table,
    ["狀態名稱", "枚舉值", "觸發條件", "行為", "可轉換目標"],
    "2E5090")

fsm_rows = [
    ("PASSIVE",         "0",  "初始化（bridge_MPC_to_RL=False 時不啟動）",
     "無力矩輸出，馬達自由旋轉",
     "任意"),
    ("RECOVERY_STAND",  "6",  "預設初始狀態；locomotionUnsafe=True；estop；roll/pitch 超限",
     "三階段：FoldLegs→RollOver→StandUp\n線性插值 PD 控制",
     "LOCOMOTION, PASSIVE"),
    ("LOCOMOTION",      "4",  "gamepad is_moving=True，UDP mode='locomotion'",
     "Convex MPC 求解 GRF，Jacobian 映射關節力矩\n安全檢查 roll<40°, pitch<40°",
     "RECOVERY_STAND, PASSIVE"),
]

for i, row_data in enumerate(fsm_rows):
    bg = "EBF3FF" if i % 2 == 0 else "FFFFFF"
    add_table_row(fsm_table, row_data, bg)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 7 – PD STAND
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "7. PD 站立控制（pd_stand_sdk2）", 1, (0x1F, 0x38, 0x64))

add_para(doc,
    "pd_stand_sdk2() 為一獨立於 FSM 的控制函式，"
    "在 gamepad.is_standing==True 時由 MPC_RUN() 直接呼叫，"
    "繞過 RobotRunnerFSM，直接輸出 PD 力矩。",
    size=10)

add_heading(doc, "7.1 目標關節角度", 2)
pd_target = doc.add_table(rows=1, cols=4)
pd_target.style = 'Table Grid'
pd_target.alignment = WD_TABLE_ALIGNMENT.CENTER
table_header_row(pd_target, ["關節", "Hip(rad)", "Thigh(rad)", "Knee(rad)"], "2E5090")

pd_rows = [
    ("STAND_TARGET（站立）", "0.0",  "0.8", "-1.6"),
    ("SIT_TARGET（坐下）",   "±0.0", "1.4", "-2.7"),
]
for i, r in enumerate(pd_rows):
    add_table_row(pd_target, r, "EBF3FF" if i%2==0 else "FFFFFF")

doc.add_paragraph()

add_heading(doc, "7.2 PD 增益參數", 2)
pd_gain = doc.add_table(rows=1, cols=4)
pd_gain.style = 'Table Grid'
pd_gain.alignment = WD_TABLE_ALIGNMENT.CENTER
table_header_row(pd_gain, ["腿部", "Kp（位置增益）", "Kd 矩陣（速度增益）", "適用腿"], "2E5090")

gain_rows = [
    ("前腿", "KP_FRONT = 50.0",  "diag(5, 5, 5)（3×3）", "leg 0,1（FL, FR）"),
    ("後腿", "KP_BACK  = 85.0",  "diag(7, 7, 7)（3×3）", "leg 2,3（RL, RR）"),
]
for i, r in enumerate(gain_rows):
    add_table_row(pd_gain, r, "EBF3FF" if i%2==0 else "FFFFFF")

doc.add_paragraph()

add_para(doc, "tanh 平滑過渡公式：", bold=True, size=10)
add_code_block(doc,
    "phase = tanh(elapsed_time / 1.2)  # 0→1，約 2.4s 達到穩定\n"
    "kp_val = kp * phase + 20 * (1 - phase)  # 從低增益平滑拉升\n"
    "smooth_target = phase * target + (1-phase) * q_start\n"
    "τ = Kp * (smooth_target - q) - Kd @ dq")

add_note(doc,
    "【潛在問題】pd_stand_sdk2() 使用模組層級（global）的 last_target / transition_start_time 等變數，"
    "與 pd_stand()（非 SDK2 版本）共用同一個 mujoco_sim_utils 模組中的全域狀態，"
    "若兩函式在同一 session 中交替呼叫（如 mujoco_sim.py 與 sdk2 版同時存在），"
    "可能因全域狀態污染導致過渡計時器錯誤。")

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 8 – GET BODY STATE BUG
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "8. get_body_state_sdk2() 語意錯誤分析", 1, (0x1F, 0x38, 0x64))

add_para(doc,
    "get_body_state_sdk2() 負責從 LowState_ 提取機體姿態供 StateEstimator 使用。"
    "現有實作存在一個語意上的錯誤：",
    size=10)

add_code_block(doc,
    "# mujoco_sim_utils.py, line 104\n"
    "Body_state['pose']['p'] = tuple(low_state.imu_state.rpy)   # ← 此行有誤！\n"
    "\n"
    "# 欄位定義要求 'p' = (x, y, z) 位置（Position），\n"
    "# 但實際填入 imu_state.rpy（Roll/Pitch/Yaw 歐拉角），\n"
    "# 導致 StateEstimator 的 position 資訊完全錯誤。")

add_para(doc,
    "StateEstimator.update() 的 else 分支並不讀取 pose.p（位置），"
    "僅讀取 pose.r（四元數）和 vel（速度），因此此錯誤在目前版本下不會直接影響 MPC 計算。"
    "但若未來擴充 StateEstimator 讀取 position 欄位，將引入難以排查的 Bug。",
    size=10)

add_note(doc,
    "【建議修正】應改為：Body_state['pose']['p'] = (0.0, 0.0, 0.27)（使用預設機體高度）"
    "或透過 high_state position 欄位提供真實位置。")

add_para(doc,
    "另外，get_body_state_sdk2() 將 imu_state.accelerometer 填入 vel.linear，"
    "而 StateEstimator.update() 則將 vel.linear 視為 vWorld（世界座標線速度）。"
    "加速度（m/s²）並非速度（m/s），此處存在物理量的混用，"
    "會影響 MPC 中速度追蹤的準確度。",
    size=10)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 9 – COMPLETE BUG LIST
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "9. 已識別潛在問題彙整", 1, (0x1F, 0x38, 0x64))

bug_table = doc.add_table(rows=1, cols=5)
bug_table.style = 'Table Grid'
bug_table.alignment = WD_TABLE_ALIGNMENT.CENTER
table_header_row(bug_table,
    ["#", "問題位置", "問題類型", "問題描述", "嚴重度"],
    "8B0000")

bugs = [
    ("B-01",
     "mujoco_sim_sdk2.py\nSimulationThread / LowCmdHandler",
     "執行緒安全",
     "LowCmdHandler 直接寫入 mj_data.ctrl，\nSimulationThread 同時讀取 mj_data 進行 mj_step，\n無 Lock 保護，可能造成 data race。",
     "高"),
    ("B-02",
     "mpc_locomotion_sdk2.py\nMPC_RUN / LowCmdWrite",
     "執行緒安全",
     "MPC_RUN()（200Hz）與 LowCmdWrite()（500Hz）\n共用 self.low_cmd 無 Lock 保護，\nCRC 計算與 tau 寫入可能同時發生。",
     "高"),
    ("B-03",
     "unitree_sdk2py_bridge.py\nPublishLowState",
     "屬性名稱錯誤\n（AttributeError 風險）",
     "__init__ 初始化 self.have_frame_sensor（無底線），\n但 PublishLowState 讀取 self.have_frame_sensor_（有底線），\n且條件判斷邏輯用 have_frame_sensor_ 代替 have_imu_，\n若感測器名稱不匹配將拋出 AttributeError。",
     "高"),
    ("B-04",
     "mujoco_sim_utils.py\nget_body_state_sdk2(), line 104",
     "語意錯誤",
     "pose['p'] 填入 imu_state.rpy（歐拉角）而非位置，\n物理量定義不符，目前不影響 MPC 但存在隱患。",
     "中"),
    ("B-05",
     "mujoco_sim_utils.py\nget_body_state_sdk2(), line 107",
     "物理量混用",
     "vel['linear'] 填入 accelerometer（加速度 m/s²）\n而非線速度（m/s），StateEstimator 將加速度\n誤認為世界座標線速度，影響 MPC 速度追蹤精度。",
     "中"),
    ("B-06",
     "mpc_locomotion_sdk2.py\nMPC_RUN 頻率 vs controller_dt",
     "頻率不一致",
     "MPC_RUN sleep=0.005s（200Hz），\n但 controller_dt=0.01s（100Hz），\nConvexMPC iterationsBetweenMPC 依 100Hz 計算，\n實際以 200Hz 觸發，導致 MPC 有效頻率為預期兩倍。",
     "中"),
    ("B-07",
     "mpc_locomotion_sdk2.py\n馬達索引交換（LowCmd 寫入）",
     "邏輯重複",
     "R↔L 索引交換邏輯分散於 get_dof_state_sdk2()、\npd_stand_sdk2() 與 MPC_RUN() 三處各自實作，\n維護不一致時易引入馬達控制錯誤。",
     "中"),
    ("B-08",
     "mujoco_sim_utils.py\nGlobal 狀態共用",
     "模組全域污染",
     "last_target、transition_start_time 等過渡狀態\n為模組層級全域變數，pd_stand 與 pd_stand_sdk2\n共用同一狀態，不同呼叫者可能互相干擾。",
     "低"),
    ("B-09",
     "teleop_client.py\n心跳包停用",
     "功能缺失",
     "heartbeat 執行緒已被注解，\n若使用者無操作 MPC_RUN() 仍持續以上次收到的速度執行，\n直到 socketrecvfrom timeout（50ms）才清零。",
     "低"),
    ("B-10",
     "unitree_sdk2py_bridge.py\nPublishHighState",
     "資料源錯誤",
     "high_state.position/velocity 讀取感測器索引 46-51，\n但實際 go2 感測器定義需對應確認；\nhigh_state 目前無訂閱者，資料未被利用。",
     "低"),
]

for i, row_data in enumerate(bugs):
    bg = "FFF0F0" if i % 2 == 0 else "FFFFFF"
    row = add_table_row(bug_table, row_data, bg)
    severity_cell = row.cells[4]
    sev = row_data[4]
    if sev == "高":
        for p in severity_cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
                r.font.bold = True
    elif sev == "中":
        for p in severity_cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(0xC0, 0x60, 0x00)
                r.font.bold = True

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 10 – EXPECTED EXECUTION RESULT
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "10. 執行流程與預期結果", 1, (0x1F, 0x38, 0x64))

add_heading(doc, "10.1 正常啟動流程", 2)

startup_steps = [
    ("步驟 1", "終端機 A",
     "python3 mujoco_sim_sdk2.py",
     "MuJoCo Viewer 視窗開啟，顯示 GO2 機器人場景（地板＋台階障礙）。"
     "DDS 初始化於 DOMAIN_ID=1, eth0。"
     "兩個 RecurrentThread 啟動：sim_lowstate（200Hz）、sim_highstate（200Hz）。"
     "機器人初始為靜止狀態（無力矩輸入）。"),
    ("步驟 2", "終端機 B",
     "python3 mpc_locomotion_sdk2.py",
     "[UDP_Reader] Listening for remote commands on port 9876...\n"
     "等待 rt/lowstate 首包... → 收到後離開等待迴圈。\n"
     "LowCmdWrite（500Hz）開始發佈 rt/lowcmd（tau 全零）。\n"
     "MPC_RUN 以 200Hz 執行，is_standing=False, is_moving=False，無力矩輸出。"),
    ("步驟 3", "終端機 C",
     "python3 teleop_client.py",
     "顯示操作說明，進入 CMD> 提示。"),
    ("步驟 4", "終端機 C",
     "CMD> stand [Enter]",
     "UDP 封包 {mode:'stand'} 發至 9876。\n"
     "UDPGamepad 設 is_standing=True，is_moving=False。\n"
     "MPC_RUN 呼叫 pd_stand_sdk2()，機器人開始以 tanh 曲線平滑站立（約 2~3 秒）。\n"
     "MuJoCo Viewer 可見 GO2 緩慢抬起四肢至站立姿態。"),
    ("步驟 5", "終端機 C",
     "CMD> move [Enter] → 進入 wsad 模式",
     "UDP {mode:'locomotion'} 發出。\n"
     "is_moving=True, is_standing=False。\n"
     "ControlFSM 觸發 RECOVERY_STAND → LOCOMOTION 轉換（transitionDone=True 立即完成）。\n"
     "ConvexMPCLocomotion 啟動 trot 步態，MPC 每 ~0.027s 求解一次 QP。\n"
     "機器人在 Viewer 中開始 trot 行走（原地踏步）。"),
    ("步驟 6", "終端機 C",
     "wsad: W 鍵",
     "vx += 0.5（上限 4.0）。\n"
     "Convex MPC 計算前進力矩，機器人在 Viewer 中向前移動。"),
    ("步驟 7", "終端機 C",
     "wsad: Z 鍵（煞車）",
     "vx=vy=wz=0，機器人減速至原地 trot。"),
    ("步驟 8", "終端機 C",
     "CMD> estop",
     "{estop:True}，vx=vy=wz=0，_mode=RECOVERY_STAND。\n"
     "ControlFSM 觸發 LOCOMOTION → RECOVERY_STAND 轉換。\n"
     "機器人停止行走，執行 RecoveryStand 動作（折疊腿→站立）。"),
]

for step, terminal, cmd, result in startup_steps:
    p = doc.add_paragraph()
    r1 = p.add_run(f"[{step} / {terminal}]  ")
    r1.font.bold = True
    r1.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    r1.font.size = Pt(10)
    r2 = p.add_run(cmd)
    r2.font.name = 'Courier New'
    r2.font.size = Pt(9)
    r2.font.color.rgb = RGBColor(0x00, 0x60, 0x00)

    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(1)
    r3 = p2.add_run(result)
    r3.font.size = Pt(9)

doc.add_paragraph()

# ─── 10.2 Expected viewer output ─────────────────────────────────────────
add_heading(doc, "10.2 MuJoCo Viewer 預期畫面", 2)

viewer_info = [
    "• 場景：平地＋右側台階（高 3.5cm ~ 21cm 漸增台階），GO2 模型居中。",
    "• 攝影機：預設 azimuth=90, elevation=-20, distance=2.0，俯視右側 45° 視角。",
    "• 站立階段：四肢緩慢向 q=[0, 0.8, -1.6] 靠攏，約 2-3 秒達到標準站立姿態。",
    "• trot 行走：四肢呈對角交替步態（FR+RL 同步，FL+RR 同步），頻率由 MPC 決定。",
    "• 緊急停止後：機器人執行 FoldLegs 動作後再 StandUp，完成 Recovery Balance。",
    "• Viewer 更新頻率 50Hz，物理仿真 200Hz，兩者無間隔感覺流暢。",
]
for line in viewer_info:
    add_para(doc, line, size=10, indent=0.5)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 11 – FREQUENCY SUMMARY
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "11. 訊號頻率彙整表", 1, (0x1F, 0x38, 0x64))

freq_table = doc.add_table(rows=1, cols=6)
freq_table.style = 'Table Grid'
freq_table.alignment = WD_TABLE_ALIGNMENT.CENTER
table_header_row(freq_table,
    ["通訊節點 / 執行緒", "所在程式", "週期", "頻率", "通訊介質", "備註"],
    "1F3864")

freq_rows = [
    ("teleop_client 按鍵觸發",       "teleop_client.py",          "不定期",  "< 10 Hz 人工",   "UDP",        "按鍵驅動，非定時"),
    ("UDPGamepad._listen()",          "mpc_locomotion_sdk2.py",    "50ms 超時","~20 Hz max",   "UDP recv",   "socket.timeout=0.05s"),
    ("MPC_RUN() 主控制迴圈",          "mpc_locomotion_sdk2.py",    "0.005s",  "200 Hz",         "記憶體",     "time.sleep(0.005)"),
    ("LowCmdWrite / writebasiccmd",   "mpc_locomotion_sdk2.py",    "0.002s",  "500 Hz",         "DDS rt/lowcmd",  "RecurrentThread"),
    ("SimulationThread / mj_step",    "mujoco_sim_sdk2.py",        "0.005s",  "200 Hz",         "MuJoCo",     "等待 Lock"),
    ("PhysicsViewerThread",           "mujoco_sim_sdk2.py",        "0.02s",   "50 Hz",          "MuJoCo Viewer","time.sleep(0.02)"),
    ("PublishLowState / sim_lowstate","mujoco_sim_sdk2.py",        "0.005s",  "200 Hz",         "DDS rt/lowstate","RecurrentThread(dt)"),
    ("PublishHighState/sim_highstate","mujoco_sim_sdk2.py",        "0.005s",  "200 Hz",         "DDS rt/sportmodestate","RecurrentThread(dt)"),
    ("LowCmdHandler (callback)",      "mujoco_sim_sdk2.py",        "事件觸發","≤500 Hz",        "DDS 回呼",   "DDS 收包即觸發"),
    ("LowStateHandler (callback)",    "mpc_locomotion_sdk2.py",    "事件觸發","≤200 Hz",        "DDS 回呼",   "DDS 收包即觸發"),
    ("Convex MPC QP 求解",            "ConvexMPCLocomotion.run()", "每2次FSM","~100 Hz trigger\n~37-50Hz MPC","CPU計算",     "iterationsBetweenMPC≈2"),
]

for i, row_data in enumerate(freq_rows):
    bg = "EBF3FF" if i % 2 == 0 else "FFFFFF"
    add_table_row(freq_table, row_data, bg)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 12 – SUGGESTIONS
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "12. 改善建議", 1, (0x1F, 0x38, 0x64))

suggestions = [
    ("S-01", "加入 threading.Lock 保護 low_cmd",
     "在 MPC_RUN() 寫入 tau 與 LowCmdWrite() 讀取/CRC 之間加入 mutex lock，"
     "避免 data race。建議使用 threading.Lock() 並在兩個執行緒的臨界區段 acquire/release。"),
    ("S-02", "修正 have_frame_sensor_ / have_imu_ 初始化",
     "在 UnitreeSdk2Bridge.__init__ 中明確初始化 self.have_frame_sensor_ = False 與 "
     "self.have_imu_ = False，並修正 PublishLowState 的判斷條件為 have_imu_。"),
    ("S-03", "修正 get_body_state_sdk2() 位置與速度欄位",
     "將 pose['p'] 改為機體位置（可用 high_state.position 或常數 body height）；"
     "將 vel['linear'] 改為真實線速度來源。"),
    ("S-04", "統一馬達索引交換邏輯",
     "將 R↔L 索引交換抽取為獨立工具函式，集中於 mujoco_sim_utils.py，"
     "由 get_dof_state_sdk2、pd_stand_sdk2、MPC_RUN 統一呼叫，避免三處各自維護。"),
    ("S-05", "對齊 controller_dt 與主迴圈週期",
     "將 MPC_RUN 的 sleep 改為 0.01s（=controller_dt）以對齊 Parameters.controller_dt，"
     "或更新 iterationsBetweenMPC 計算以反映實際 200Hz 觸發頻率。"),
    ("S-06", "重新啟用或重設計心跳包機制",
     "為 teleop_client 加入定時心跳包（建議 0.05s），"
     "MPC_RUN 端增加超時偵測（如 0.5s 無封包則自動 estop），"
     "防止通訊中斷後機器人繼續前進。"),
    ("S-07", "新增 SimulationThread 與 LowCmdHandler 的 mj_data.ctrl 保護",
     "在 LowCmdHandler 寫入 mj_data.ctrl 時也應持有 locker，"
     "確保 mj_step 讀取 ctrl 時不被同時寫入。"),
]

for sid, title, detail in suggestions:
    p = doc.add_paragraph()
    r1 = p.add_run(f"[{sid}] {title}")
    r1.font.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(1)
    r2 = p2.add_run(detail)
    r2.font.size = Pt(9.5)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 13 – CONFIG SUMMARY
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "13. 全域設定參數彙整（config_sdk2.py & Parameters.py）", 1, (0x1F, 0x38, 0x64))

cfg_table = doc.add_table(rows=1, cols=4)
cfg_table.style = 'Table Grid'
cfg_table.alignment = WD_TABLE_ALIGNMENT.CENTER
table_header_row(cfg_table, ["參數名稱", "值", "影響範圍", "說明"], "2E5090")

cfg_rows = [
    ("DOMAIN_ID",           "1",        "DDS 通訊",        "CycloneDDS 域 ID，兩端需一致"),
    ("INTERFACE",           "eth0",     "DDS 通訊",        "無 multicast 環境下指定介面"),
    ("SIMULATE_DT",         "0.005s",   "物理仿真/DDS發佈","MuJoCo timestep，200Hz"),
    ("VIEWER_DT",           "0.02s",    "Viewer 渲染",     "50Hz，不可低於 viewer.sync() 執行時間"),
    ("PRINT_SCENE_INFORMATION","0",     "除錯",            "設 1 可列印 Link/Joint/Sensor 資訊"),
    ("Parameters.controller_dt","0.01s","MPC 計算",       "FSM 預期控制週期 100Hz"),
    ("Parameters.cmpc_gait","TROT",     "步態選擇",        "預設 Trot，可由 UDP 指令切換"),
    ("Parameters.control_mode","RECOVERY_STAND","FSM初始狀態","啟動後先進 Recovery Stand"),
    ("Parameters.locomotionUnsafe","False","安全旗標",     "True 時 FSM 觸發返回 Recovery Stand"),
    ("Parameters.FSM_check_safety","True","安全檢查",      "roll/pitch 超 40° 自動切回 Recovery"),
    ("KP_FRONT",            "50.0",     "PD 站立-前腿",    "位置增益"),
    ("KP_BACK",             "85.0",     "PD 站立-後腿",    "後腿增益較高維持穩定"),
    ("KD_FRONT",            "diag(5,5,5)","PD 站立-前腿", "速度阻尼"),
    ("KD_BACK",             "diag(7,7,7)","PD 站立-後腿", "速度阻尼"),
    ("STAND_TARGET",        "[0,0.8,-1.6]×4","PD 站立目標","Hip/Thigh/Knee 目標角度（rad）"),
    ("SIT_TARGET",          "[0,1.4,-2.7]×4","PD 坐下目標","目前未透過 SDK2 介面動態切換"),
]

for i, row_data in enumerate(cfg_rows):
    bg = "EBF3FF" if i % 2 == 0 else "FFFFFF"
    add_table_row(cfg_table, row_data, bg)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  FOOTER / END
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()

end_p = doc.add_paragraph()
end_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
end_r = end_p.add_run("— 報告書結束 —")
end_r.font.size = Pt(11)
end_r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
end_r.font.bold = True

doc.add_paragraph()

note_p = doc.add_paragraph()
note_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
note_r = note_p.add_run(
    "本文件由自動化程式碼分析工具生成\n"
    "基於 branch: pd_sit2stand | 分析日期: 2026-05-10\n"
    "覆蓋程式：mpc_locomotion_sdk2.py · mujoco_sim_sdk2.py · teleop_client.py\n"
    "         unitree_sdk2py_bridge.py · udp_reader.py · mujoco_sim_utils.py · config_sdk2.py\n"
    "         MPC_Controller/** (RobotRunnerFSM, ControlFSM, LegController, StateEstimator,\n"
    "                            ConvexMPCLocomotion, FSM_State_*, Quadruped, Parameters)"
)
note_r.font.size = Pt(8)
note_r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

# ─── Save ──────────────────────────────────────────────────────────────────
output_path = "/home/ScarpGhost/mujoco_sim/docs/SDK2_Architecture_Analysis_Report.docx"
doc.save(output_path)
print(f"Report saved to: {output_path}")
